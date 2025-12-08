"""
DOCX semantic extractor for golden tests.

Extracts meaningful content from DOCX files into a normalized JSON structure,
avoiding brittle binary comparisons.

Extracted elements:
1. Body paragraphs (text content)
2. Table cells (table index, row, column, text)
3. Image metadata (count, relationship IDs, dimensions if available)
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from zipfile import ZipFile

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.table import Table
from docx.text.paragraph import Paragraph

from .normalize import Normalizer


@dataclass
class ImageInfo:
    """Extracted image information."""
    rel_id: str
    filename: str | None = None
    width: int | None = None
    height: int | None = None
    content_type: str | None = None
    alt_text: str | None = None
    slot_key: str | None = None  # Inferred from filename pattern


@dataclass
class TableCell:
    """Extracted table cell."""
    table_index: int
    row: int
    col: int
    text: str


@dataclass
class DocxContent:
    """Extracted DOCX content structure."""
    paragraphs: list[str] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)
    images: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        """Convert to dictionary for JSON serialization."""
        return {
            "paragraphs": self.paragraphs,
            "tables": self.tables,
            "images": self.images,
            "metadata": self.metadata,
        }


class DocxExtractor:
    """
    Extract semantic content from DOCX files.

    Usage:
        extractor = DocxExtractor()
        content = extractor.extract(Path("document.docx"))
        json_data = content.to_dict()
    """

    def __init__(
        self,
        normalizer: Normalizer | None = None,
        include_empty_paragraphs: bool = False,
        include_image_details: bool = True,
    ):
        """
        Args:
            normalizer: Normalizer instance for text processing
            include_empty_paragraphs: Whether to include empty paragraphs
            include_image_details: Whether to extract detailed image info
        """
        self.normalizer = normalizer or Normalizer()
        self.include_empty_paragraphs = include_empty_paragraphs
        self.include_image_details = include_image_details

    def extract(self, docx_path: Path) -> DocxContent:
        """
        Extract content from a DOCX file.

        Args:
            docx_path: Path to the DOCX file

        Returns:
            DocxContent with extracted elements
        """
        doc = Document(docx_path)
        content = DocxContent()

        # Extract paragraphs
        content.paragraphs = self._extract_paragraphs(doc)

        # Extract tables
        content.tables = self._extract_tables(doc)

        # Extract images
        content.images = self._extract_images(doc, docx_path)

        # Extract metadata
        content.metadata = self._extract_metadata(doc)

        return content

    def _extract_paragraphs(self, doc: Document) -> list[str]:
        """Extract body paragraph texts."""
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text
            normalized = self.normalizer.normalize(text)

            if normalized or self.include_empty_paragraphs:
                paragraphs.append(normalized if normalized else "")

        return paragraphs

    def _extract_tables(self, doc: Document) -> list[list[list[str]]]:
        """
        Extract table contents.

        Returns a list of tables, where each table is a 2D list of cell texts.
        """
        tables = []

        for table in doc.tables:
            table_data = []

            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Get cell text and normalize
                    cell_text = cell.text
                    normalized = self.normalizer.normalize(cell_text)
                    row_data.append(normalized if normalized else "")
                table_data.append(row_data)

            tables.append(table_data)

        return tables

    def _extract_images(
        self,
        doc: Document,
        docx_path: Path,
    ) -> list[dict[str, Any]]:
        """
        Extract image information.

        Returns list of image metadata (not the actual binary data).
        Includes:
        - Relationship count (how many image relationships exist)
        - Media file count (how many files in word/media/)
        - Per-image details (filename, size, inferred slot)
        """
        images = []
        relationship_count = 0
        media_file_count = 0

        # Get all image relationships from the document
        try:
            # Access the document's relationship parts
            rels = doc.part.rels

            for rel_id, rel in rels.items():
                if rel.reltype == RT.IMAGE:
                    relationship_count += 1

                    image_info: dict[str, Any] = {
                        "rel_id": rel_id,
                        "count_position": len(images) + 1,
                    }

                    if self.include_image_details:
                        # Get the target filename
                        filename = None
                        if hasattr(rel, 'target_ref'):
                            # target_ref is the actual path string
                            target = str(rel.target_ref)
                            filename = target.split('/')[-1]
                            image_info["filename"] = filename
                        elif hasattr(rel, '_target'):
                            # Fallback: _target might be a string or object
                            target = rel._target
                            if isinstance(target, str):
                                filename = target.split('/')[-1]
                            else:
                                # It's an object, try to get partname
                                try:
                                    filename = target.partname.split('/')[-1]
                                except Exception:
                                    filename = None
                            if filename:
                                image_info["filename"] = filename

                        # Infer slot from filename pattern
                        if filename:
                            image_info["inferred_slot"] = self._infer_slot_from_filename(filename)

                        # Try to get image dimensions from the blob
                        try:
                            blob = rel.target_part.blob
                            image_info["size_bytes"] = len(blob)
                        except Exception:
                            pass

                    images.append(image_info)

        except Exception:
            pass

        # Also count media files from ZIP structure (more reliable count)
        try:
            with ZipFile(docx_path) as zf:
                media_files = [
                    f for f in zf.namelist()
                    if f.startswith('word/media/')
                ]
                media_file_count = len(media_files)

                # If we didn't get relationships, use ZIP fallback
                if not images:
                    for i, media_file in enumerate(media_files):
                        filename = media_file.split('/')[-1]
                        images.append({
                            "rel_id": f"media_{i+1}",
                            "filename": filename,
                            "count_position": i + 1,
                            "inferred_slot": self._infer_slot_from_filename(filename),
                        })
        except Exception:
            pass

        # Add summary counts for verification
        if images:
            # Add counts to first image for easy access in tests
            images[0]["_image_summary"] = {
                "total_count": len(images),
                "relationship_count": relationship_count,
                "media_file_count": media_file_count,
            }

        return images

    def _infer_slot_from_filename(self, filename: str) -> str | None:
        """
        Infer photo slot from image filename.

        Examples:
            image1.jpeg → None (generic)
            overview.jpg → "overview"
            label_serial.png → "label_serial"
        """
        if not filename:
            return None

        # Remove extension
        stem = filename.rsplit('.', 1)[0] if '.' in filename else filename

        # Known slot patterns
        slot_patterns = [
            "overview",
            "label_serial",
            "measurement_setup",
            "defect",
            "detail",
        ]

        # Check for exact or partial match
        stem_lower = stem.lower()
        for slot in slot_patterns:
            if slot in stem_lower:
                return slot

        # Generic image names (image1, image2, etc.) → None
        if re.match(r'^image\d+$', stem_lower):
            return None

        return None

    def _extract_metadata(self, doc: Document) -> dict[str, Any]:
        """Extract document metadata (normalized)."""
        metadata: dict[str, Any] = {}

        try:
            core_props = doc.core_properties

            # Only include stable metadata (not dates/times)
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.subject:
                metadata["subject"] = core_props.subject

        except Exception:
            pass

        return metadata

    def extract_to_dict(self, docx_path: Path) -> dict[str, Any]:
        """
        Extract and return as dictionary.

        Convenience method for JSON serialization.
        """
        content = self.extract(docx_path)
        return content.to_dict()


def extract_docx(
    docx_path: Path,
    normalizer: Normalizer | None = None,
) -> dict[str, Any]:
    """
    Convenience function to extract DOCX content.

    Args:
        docx_path: Path to DOCX file
        normalizer: Optional normalizer instance

    Returns:
        Dictionary of extracted content
    """
    extractor = DocxExtractor(normalizer=normalizer)
    return extractor.extract_to_dict(docx_path)
