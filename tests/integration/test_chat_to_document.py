"""
test_chat_to_document.py - 채팅 → 문서 생성 통합 테스트

검증 포인트:
- 채팅 입력 → AI 파싱 (mock) → 문서 생성
- intake_session.json append-only
- AI Provider mock
"""

import json
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
import yaml
from docx import Document

from src.app.services.intake import IntakeService
from src.app.services.extract import ExtractionService
from src.app.services.validate import ValidationService
from src.app.providers.base import OCRResult, ExtractionResult
from src.core.ssot_job import ensure_job_json
from src.core.ids import generate_job_id
from src.render.word import DocxRenderer


# =============================================================================
# Fixtures
# =============================================================================

@pytest.fixture
def chat_root(tmp_path: Path) -> Path:
    """채팅 테스트 루트 디렉터리."""
    root = tmp_path / "chat_test"
    root.mkdir()
    return root


@pytest.fixture
def definition_path(chat_root: Path) -> Path:
    """definition.yaml 생성."""
    definition = {
        "definition_version": "1.0.0",
        "fields": {
            "wo_no": {
                "type": "token",
                "importance": "critical",
                "aliases": ["WO No", "작업번호", "W/O"],
            },
            "line": {
                "type": "token",
                "importance": "critical",
                "aliases": ["Line", "라인"],
            },
            "part_no": {
                "type": "token",
                "importance": "critical",
                "aliases": ["Part No", "품번"],
            },
            "lot": {
                "type": "token",
                "importance": "critical",
                "aliases": ["LOT"],
            },
            "result": {
                "type": "token",
                "importance": "critical",
                "aliases": ["Result", "결과", "판정"],
            },
            "inspector": {
                "type": "token",
                "importance": "reference",
                "aliases": ["Inspector", "검사자"],
            },
        },
        "validation": {
            "result_pass_aliases": ["PASS", "OK", "합격"],
            "result_fail_aliases": ["FAIL", "NG", "불합격"],
        },
    }

    path = chat_root / "definition.yaml"
    with open(path, "w", encoding="utf-8") as f:
        yaml.dump(definition, f, allow_unicode=True)

    return path


@pytest.fixture
def prompts_dir(chat_root: Path) -> Path:
    """프롬프트 디렉터리."""
    prompts = chat_root / "prompts"
    prompts.mkdir()
    return prompts


@pytest.fixture
def docx_template(chat_root: Path) -> Path:
    """테스트용 DOCX 템플릿."""
    template_path = chat_root / "template.docx"

    doc = Document()
    doc.add_heading("검사 성적서", 0)
    doc.add_paragraph("작업번호: {{wo_no}}")
    doc.add_paragraph("라인: {{line}}")
    doc.add_paragraph("결과: {{result}}")
    doc.save(template_path)

    return template_path


@pytest.fixture
def mock_llm_provider():
    """Mock LLM Provider."""
    provider = MagicMock()
    provider.extract_fields = AsyncMock(return_value=ExtractionResult(
        success=True,
        fields={
            "wo_no": "WO-001",
            "line": "L1",
            "part_no": "PART-A",
            "lot": "LOT-001",
            "result": "PASS",
            "inspector": "홍길동",
        },
        model_requested="claude-opus-4-5-20251101",
        model_used="claude-opus-4-5-20251101",
    ))
    return provider


@pytest.fixture
def mock_ocr_result():
    """Mock OCR Result."""
    return OCRResult(
        success=True,
        text="""
        검사 성적서

        작업번호: WO-001
        라인: L1
        품번: PART-A
        LOT: LOT-001
        결과: PASS
        검사자: 홍길동
        """,
        confidence=0.95,
        model_requested="gemini-3-pro",
        model_used="gemini-3-pro",
    )


# =============================================================================
# 1. IntakeService 통합 테스트
# =============================================================================

class TestIntakeServiceIntegration:
    """IntakeService 통합 테스트."""

    def test_full_intake_flow(self, chat_root, mock_ocr_result):
        """전체 Intake 흐름."""
        job_dir = chat_root / "jobs" / "test_job"
        service = IntakeService(job_dir)

        # 1. 세션 생성
        session = service.create_session()
        assert session.session_id is not None

        # 2. 사용자 메시지 추가
        msg1 = service.add_message("user", "작업번호 WO-001, 라인 L1입니다")
        assert msg1.role == "user"

        # 3. 어시스턴트 응답 추가
        msg2 = service.add_message("assistant", "WO-001 확인했습니다")
        assert msg2.role == "assistant"

        # 4. OCR 결과 추가
        service.add_ocr_result("image.jpg", mock_ocr_result)

        # 5. 세션 확인
        loaded = service.load_session()
        assert len(loaded.messages) == 2
        assert "image.jpg" in loaded.ocr_results

    def test_intake_append_only(self, chat_root):
        """Intake append-only 검증."""
        job_dir = chat_root / "jobs" / "test_job"
        service = IntakeService(job_dir)

        service.create_session()
        service.add_message("user", "Message 1")
        service.add_message("user", "Message 2")
        service.add_message("user", "Message 3")

        session = service.load_session()

        # 메시지가 순서대로 append됨
        assert len(session.messages) == 3
        assert session.messages[0].content == "Message 1"
        assert session.messages[1].content == "Message 2"
        assert session.messages[2].content == "Message 3"


# =============================================================================
# 2. Extraction 통합 테스트
# =============================================================================

class TestExtractionIntegration:
    """Extraction 서비스 통합 테스트."""

    @pytest.mark.asyncio
    async def test_regex_extraction(self, definition_path, prompts_dir, mock_llm_provider):
        """정규식 추출 테스트."""
        config = {"ai": {"llm": {"model": "claude-opus-4-5-20251101"}}}

        service = ExtractionService(
            config=config,
            definition_path=definition_path,
            prompts_dir=prompts_dir,
            provider=mock_llm_provider,
        )

        # 모든 critical 필드가 정규식으로 추출 가능한 입력
        user_input = """
        작업번호: WO-001
        라인: L1
        품번: PART-A
        LOT: LOT-001
        결과: PASS
        """

        result = await service.extract(user_input)

        # 정규식만으로 추출됨 (LLM 호출 안 함)
        assert result.model_used == "regex"
        assert result.fields["wo_no"] == "WO-001"
        mock_llm_provider.extract_fields.assert_not_called()

    @pytest.mark.asyncio
    async def test_llm_extraction_fallback(
        self,
        definition_path,
        prompts_dir,
        mock_llm_provider,
    ):
        """정규식 실패 시 LLM 호출."""
        config = {"ai": {"llm": {"model": "claude-opus-4-5-20251101"}}}

        service = ExtractionService(
            config=config,
            definition_path=definition_path,
            prompts_dir=prompts_dir,
            provider=mock_llm_provider,
        )

        # 불완전한 입력 (정규식으로 모든 필드 추출 불가)
        user_input = "작업번호 WO-001입니다"

        result = await service.extract(user_input)

        # LLM 호출됨
        mock_llm_provider.extract_fields.assert_called_once()

    @pytest.mark.asyncio
    async def test_regex_and_llm_merge(
        self,
        definition_path,
        prompts_dir,
    ):
        """정규식 + LLM 결과 병합."""
        mock_provider = MagicMock()
        mock_provider.extract_fields = AsyncMock(return_value=ExtractionResult(
            success=True,
            fields={
                "wo_no": "WO-FROM-LLM",  # LLM이 다른 값 반환
                "line": "L1",
                "part_no": "PART-A",
                "lot": "LOT-001",
                "result": "PASS",
            },
            model_used="claude",
        ))

        config = {"ai": {"llm": {"model": "claude-opus-4-5-20251101"}}}

        service = ExtractionService(
            config=config,
            definition_path=definition_path,
            prompts_dir=prompts_dir,
            provider=mock_provider,
        )

        # 정규식으로 wo_no만 추출 가능
        user_input = "작업번호: WO-001"

        result = await service.extract(user_input)

        # 정규식 결과가 우선
        assert result.fields["wo_no"] == "WO-001"
        # LLM 결과로 나머지 채움
        assert result.fields["line"] == "L1"


# =============================================================================
# 3. 채팅 → 문서 전체 흐름 테스트
# =============================================================================

class TestChatToDocumentFlow:
    """채팅 → 문서 생성 전체 흐름."""

    @pytest.mark.asyncio
    async def test_chat_input_to_document(
        self,
        chat_root,
        definition_path,
        prompts_dir,
        docx_template,
        mock_llm_provider,
    ):
        """채팅 입력 → 문서 생성 전체 흐름."""
        jobs_dir = chat_root / "jobs"

        # 1. Intake: 사용자 입력 수집
        job_dir = jobs_dir / "test_job"
        intake_service = IntakeService(job_dir)
        intake_service.create_session()

        user_input = "작업번호 WO-001, 라인 L1, 품번 PART-A, LOT LOT-001, 결과 합격"
        intake_service.add_message("user", user_input)

        # 2. Extraction: 필드 추출
        config = {"ai": {"llm": {"model": "claude-opus-4-5-20251101"}}}
        extract_service = ExtractionService(
            config=config,
            definition_path=definition_path,
            prompts_dir=prompts_dir,
            provider=mock_llm_provider,
        )

        extraction_result = await extract_service.extract(user_input)

        # Intake에 추출 결과 저장
        intake_service.add_extraction_result(ExtractionResult(
            success=True,
            fields=extraction_result.fields,
            model_requested="claude-opus-4-5-20251101",
            model_used=extraction_result.model_used,
        ))

        # 3. Validation: 검증
        validate_service = ValidationService(definition_path)
        final_fields = intake_service.get_final_fields()
        validation_result = validate_service.validate(final_fields)

        # 정규식으로 필수 필드 모두 추출됨
        assert validation_result.valid is True

        # 4. SSOT: Job 생성
        job_id = generate_job_id(final_fields)
        config = {"lock": {"timeout": 5}}
        ensure_job_json(job_dir, final_fields, config, generate_job_id)

        # 5. Render: 문서 생성
        output_path = job_dir / "deliverables" / "report.docx"
        renderer = DocxRenderer(docx_template)
        renderer.render(final_fields, output_path)

        # 검증
        assert output_path.exists()

        doc = Document(output_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "WO-001" in full_text
        assert "L1" in full_text

    @pytest.mark.asyncio
    async def test_chat_with_user_correction(
        self,
        chat_root,
        definition_path,
        prompts_dir,
        docx_template,
        mock_llm_provider,
    ):
        """사용자 수정 반영 테스트."""
        job_dir = chat_root / "jobs" / "correction_test"

        # 1. Intake 세션
        intake_service = IntakeService(job_dir)
        intake_service.create_session()

        # 2. 추출 결과 저장
        intake_service.add_extraction_result(ExtractionResult(
            success=True,
            fields={
                "wo_no": "WO-001",
                "line": "L1",
                "part_no": "PART-A",
                "lot": "LOT-001",
                "result": "PASS",
            },
            model_used="regex",
        ))

        # 3. 사용자 수정
        intake_service.add_user_correction(
            field="wo_no",
            original="WO-001",
            corrected="WO-001A",
            user="operator",
        )

        # 4. 최종 필드 확인
        final_fields = intake_service.get_final_fields()

        assert final_fields["wo_no"] == "WO-001A"  # 수정 반영
        assert final_fields["line"] == "L1"  # 원본 유지

    @pytest.mark.asyncio
    async def test_ocr_to_extraction_flow(
        self,
        chat_root,
        definition_path,
        prompts_dir,
        mock_llm_provider,
        mock_ocr_result,
    ):
        """OCR → Extraction 흐름."""
        job_dir = chat_root / "jobs" / "ocr_test"

        # 1. Intake
        intake_service = IntakeService(job_dir)
        intake_service.create_session()

        # 2. 파일 업로드 (시뮬레이션)
        intake_service.add_message(
            "user",
            "사진 첨부합니다",
            attachments=[("image.jpg", b"fake image data")],
        )

        # 3. OCR 결과 저장
        intake_service.add_ocr_result("image.jpg", mock_ocr_result)

        # 4. Extraction (OCR 텍스트 사용)
        config = {"ai": {"llm": {"model": "claude-opus-4-5-20251101"}}}
        extract_service = ExtractionService(
            config=config,
            definition_path=definition_path,
            prompts_dir=prompts_dir,
            provider=mock_llm_provider,
        )

        result = await extract_service.extract(
            user_input="이 이미지 분석해주세요",
            ocr_text=mock_ocr_result.text,
        )

        # OCR 텍스트에서 정규식으로 추출
        assert result.fields["wo_no"] == "WO-001"


# =============================================================================
# 4. 세션 불변성 테스트
# =============================================================================

class TestSessionImmutability:
    """세션 불변성 테스트."""

    def test_extraction_result_cannot_be_overwritten(self, chat_root):
        """extraction_result 덮어쓰기 불가."""
        from src.domain.errors import PolicyRejectError

        job_dir = chat_root / "jobs" / "immutable_test"
        intake_service = IntakeService(job_dir)
        intake_service.create_session()

        # 첫 번째 추출 결과
        intake_service.add_extraction_result(ExtractionResult(
            success=True,
            fields={"wo_no": "WO-001"},
            model_used="regex",
        ))

        # 두 번째 추출 시도 → 에러
        with pytest.raises(PolicyRejectError):
            intake_service.add_extraction_result(ExtractionResult(
                success=True,
                fields={"wo_no": "WO-002"},
                model_used="regex",
            ))

    def test_user_corrections_are_appended(self, chat_root):
        """사용자 수정은 append만 가능."""
        job_dir = chat_root / "jobs" / "corrections_test"
        intake_service = IntakeService(job_dir)
        intake_service.create_session()

        intake_service.add_extraction_result(ExtractionResult(
            success=True,
            fields={"wo_no": "WO-001"},
            model_used="regex",
        ))

        # 여러 번 수정 가능 (히스토리 유지)
        intake_service.add_user_correction("wo_no", "WO-001", "WO-001A")
        intake_service.add_user_correction("wo_no", "WO-001A", "WO-001B")

        session = intake_service.load_session()

        assert len(session.user_corrections) == 2
        assert session.user_corrections[0].original == "WO-001"
        assert session.user_corrections[1].original == "WO-001A"
