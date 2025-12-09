"""
test_pipeline_flow.py - 전체 파이프라인 흐름 통합 테스트

검증 포인트:
- Ingest → Validate → SSOT → Photos → Hash → Render
- job.json 생성 및 SSOT 유지
- 최종 산출물 (DOCX/XLSX) 생성
"""

from pathlib import Path

import pytest
import yaml
from docx import Document
from openpyxl import Workbook
from openpyxl.workbook.defined_name import DefinedName

from src.app.services.validate import ValidationService
from src.core.hashing import compute_packet_hash
from src.core.ids import generate_job_id
from src.core.logging import complete_run_log, create_run_log, save_run_log
from src.core.photos import safe_move, select_photo_for_slot
from src.core.ssot_job import (
    ensure_job_json,
    load_job_json,
)
from src.render.excel import ExcelRenderer
from src.render.word import DocxRenderer

# =============================================================================
# Fixtures
# =============================================================================


@pytest.fixture
def pipeline_root(tmp_path: Path) -> Path:
    """파이프라인 루트 디렉터리 설정."""
    root = tmp_path / "pipeline"
    root.mkdir()

    # jobs/ 폴더
    (root / "jobs").mkdir()

    # templates/ 폴더 구조
    templates_dir = root / "templates"
    (templates_dir / "base").mkdir(parents=True)
    (templates_dir / "custom").mkdir(parents=True)

    return root


@pytest.fixture
def definition_path(pipeline_root: Path) -> Path:
    """definition.yaml 생성."""
    definition = {
        "definition_version": "1.0.0",
        "fields": {
            "wo_no": {
                "type": "token",
                "importance": "critical",
                "aliases": ["WO No", "작업번호"],
            },
            "line": {
                "type": "token",
                "importance": "critical",
                "aliases": ["Line", "라인"],
            },
            "part_no": {
                "type": "token",
                "importance": "critical",
                "aliases": ["Part No", "품번"],
            },
            "lot": {
                "type": "token",
                "importance": "critical",
                "aliases": ["LOT"],
            },
            "result": {
                "type": "token",
                "importance": "critical",
                "aliases": ["Result", "결과"],
            },
            "inspector": {
                "type": "token",
                "importance": "reference",
                "aliases": ["Inspector", "검사자"],
            },
        },
        "validation": {
            "result_pass_aliases": ["PASS", "OK", "합격"],
            "result_fail_aliases": ["FAIL", "NG", "불합격"],
        },
    }

    path = pipeline_root / "definition.yaml"
    with open(path, "w", encoding="utf-8") as f:
        yaml.dump(definition, f, allow_unicode=True)

    return path


@pytest.fixture
def docx_template(pipeline_root: Path) -> Path:
    """테스트용 DOCX 템플릿 생성."""
    template_path = pipeline_root / "templates" / "base" / "template.docx"

    doc = Document()
    doc.add_heading("검사 성적서", 0)
    doc.add_paragraph("작업번호: {{wo_no}}")
    doc.add_paragraph("라인: {{line}}")
    doc.add_paragraph("품번: {{part_no}}")
    doc.add_paragraph("LOT: {{lot}}")
    doc.add_paragraph("결과: {{result}}")
    doc.add_paragraph("검사자: {{inspector}}")
    doc.save(template_path)

    return template_path


@pytest.fixture
def xlsx_template(pipeline_root: Path) -> Path:
    """테스트용 XLSX 템플릿 생성."""
    template_path = pipeline_root / "templates" / "base" / "template.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "검사"

    ws["A1"] = "작업번호"
    ws["A2"] = "라인"
    ws["A3"] = "결과"

    # Named Ranges
    wb.defined_names.add(DefinedName("WO_NO", attr_text="검사!$B$1"))
    wb.defined_names.add(DefinedName("LINE", attr_text="검사!$B$2"))
    wb.defined_names.add(DefinedName("RESULT", attr_text="검사!$B$3"))

    wb.save(template_path)

    return template_path


@pytest.fixture
def sample_packet() -> dict:
    """테스트용 패킷 데이터."""
    return {
        "wo_no": "WO-001",
        "line": "L1",
        "part_no": "PART-A",
        "lot": "LOT-001",
        "result": "PASS",
        "inspector": "홍길동",
    }


@pytest.fixture
def sample_photos(pipeline_root: Path) -> dict[str, Path]:
    """테스트용 사진 파일 생성."""
    photos_dir = pipeline_root / "incoming_photos"
    photos_dir.mkdir()

    photos = {}
    for name in ["overview.jpg", "label.jpg", "detail.jpg"]:
        photo_path = photos_dir / name
        photo_path.write_bytes(b"fake image data for " + name.encode())
        photos[name.split(".")[0]] = photo_path

    return photos


# =============================================================================
# 1. Job 생성 및 SSOT 테스트
# =============================================================================


class TestJobCreationFlow:
    """Job 생성 흐름 테스트."""

    def test_generate_job_id_format(self, sample_packet):
        """job_id 생성 형식."""
        # generate_job_id takes a packet dict
        job_id = generate_job_id(sample_packet)

        # Job ID는 WO_001 형식으로 변환됨 (하이픈 → 언더스코어)
        assert "WO_001" in job_id or "WO-001" in job_id or "WO001" in job_id
        assert "L1" in job_id

    def test_ensure_job_json_creates_structure(self, pipeline_root, sample_packet):
        """job.json 및 폴더 구조 생성."""
        jobs_dir = pipeline_root / "jobs"

        job_id = generate_job_id(sample_packet)
        job_dir = jobs_dir / job_id
        job_dir.mkdir(parents=True)

        # 새 시그니처: (job_dir, packet, config, generate_job_id_func)
        config = {"lock": {"timeout": 5}}
        ensure_job_json(
            job_dir=job_dir,
            packet=sample_packet,
            config=config,
            generate_job_id_func=generate_job_id,
        )

        # 폴더 구조 확인
        assert (job_dir / "job.json").exists()

    def test_job_json_ssot_maintained(self, pipeline_root, sample_packet):
        """job.json SSOT 유지."""
        jobs_dir = pipeline_root / "jobs"
        job_id = generate_job_id(sample_packet)
        job_dir = jobs_dir / job_id
        job_dir.mkdir(parents=True)

        config = {"lock": {"timeout": 5}}

        # 첫 번째 생성
        ensure_job_json(job_dir, sample_packet, config, generate_job_id)
        job1 = load_job_json(job_dir / "job.json")

        # 두 번째 호출 (동일 데이터)
        ensure_job_json(job_dir, sample_packet, config, generate_job_id)
        job2 = load_job_json(job_dir / "job.json")

        # 동일 job_id 유지
        assert job1["job_id"] == job2["job_id"]


# =============================================================================
# 2. Validation 흐름 테스트
# =============================================================================


class TestValidationFlow:
    """검증 흐름 테스트."""

    def test_validates_complete_packet(self, definition_path, sample_packet):
        """완전한 패킷 검증 통과."""
        service = ValidationService(definition_path)

        result = service.validate(sample_packet)

        assert result.valid is True
        assert result.missing_required == []

    def test_normalizes_result_field(self, definition_path, sample_packet):
        """result 필드 정규화."""
        service = ValidationService(definition_path)
        sample_packet["result"] = "합격"

        result = service.validate(sample_packet)

        assert result.valid is True
        assert result.fields["result"] == "PASS"

    def test_detects_missing_critical(self, definition_path):
        """critical 필드 누락 감지."""
        service = ValidationService(definition_path)
        incomplete = {
            "wo_no": "WO-001",
            # line, part_no, lot, result 누락
        }

        result = service.validate(incomplete)

        assert result.valid is False


# =============================================================================
# 3. Photos 흐름 테스트
# =============================================================================


class TestPhotosFlow:
    """사진 처리 흐름 테스트."""

    def test_safe_move_to_photos_dir(self, pipeline_root, sample_photos):
        """사진 파일 안전 이동."""
        job_dir = pipeline_root / "jobs" / "test_job"
        photos_dir = job_dir / "photos"
        photos_dir.mkdir(parents=True)

        src = sample_photos["overview"]
        dst = photos_dir / "01_overview.jpg"

        result = safe_move(src, dst)

        assert result.success is True
        assert dst.exists()
        assert not src.exists()

    def test_photo_slot_selection(self, pipeline_root, definition_path, sample_photos):
        """사진 슬롯 선택."""
        from src.domain.schemas import PhotoSlot

        raw_dir = pipeline_root / "jobs" / "test_job" / "photos" / "raw"
        raw_dir.mkdir(parents=True)

        # 사진 복사
        src = sample_photos["overview"]
        dst = raw_dir / "01_overview.jpg"
        dst.write_bytes(src.read_bytes())

        # PhotoSlot 시그니처: key, basename, required, override_allowed, path
        slot = PhotoSlot(
            key="overview",
            basename="01_overview",
            required=True,
            override_allowed=False,
        )

        result_path, warning = select_photo_for_slot(slot, raw_dir, definition_path)

        assert result_path is not None
        assert result_path.name == "01_overview.jpg"


# =============================================================================
# 4. Hashing 흐름 테스트
# =============================================================================


class TestHashingFlow:
    """해시 생성 흐름 테스트."""

    def test_packet_hash_reproducibility(self, sample_packet, definition_path):
        """패킷 해시 재현성."""
        config = {
            "pipeline": {"hash_fields": ["wo_no", "line", "part_no", "lot", "result"]}
        }

        hash1 = compute_packet_hash(sample_packet, config, definition_path)
        hash2 = compute_packet_hash(sample_packet, config, definition_path)

        assert hash1 == hash2

    def test_packet_hash_changes_on_data_change(self, sample_packet, definition_path):
        """데이터 변경 시 해시 변경."""
        config = {
            "pipeline": {"hash_fields": ["wo_no", "line", "part_no", "lot", "result"]}
        }

        hash1 = compute_packet_hash(sample_packet, config, definition_path)

        modified = dict(sample_packet)
        modified["wo_no"] = "WO-002"
        hash2 = compute_packet_hash(modified, config, definition_path)

        assert hash1 != hash2


# =============================================================================
# 5. Render 흐름 테스트
# =============================================================================


class TestRenderFlow:
    """렌더링 흐름 테스트."""

    def test_render_docx(self, docx_template, sample_packet, tmp_path):
        """DOCX 렌더링."""
        output_path = tmp_path / "output.docx"

        renderer = DocxRenderer(docx_template)
        result = renderer.render(sample_packet, output_path)

        assert result.exists()

        # 내용 확인
        doc = Document(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "WO-001" in full_text
        assert "L1" in full_text
        assert "PASS" in full_text

    def test_render_xlsx(self, xlsx_template, sample_packet, tmp_path):
        """XLSX 렌더링."""
        output_path = tmp_path / "output.xlsx"

        manifest = {
            "xlsx_mappings": {
                "named_ranges": {
                    "wo_no": "WO_NO",
                    "line": "LINE",
                    "result": "RESULT",
                },
                "cell_addresses": {},
            },
        }

        renderer = ExcelRenderer(xlsx_template, manifest)
        result = renderer.render(sample_packet, output_path)

        assert result.exists()

        # 내용 확인
        from openpyxl import load_workbook

        wb = load_workbook(result)
        ws = wb["검사"]

        assert ws["B1"].value == "WO-001"
        assert ws["B2"].value == "L1"
        assert ws["B3"].value == "PASS"


# =============================================================================
# 6. 전체 파이프라인 흐름 테스트
# =============================================================================


class TestFullPipelineFlow:
    """전체 파이프라인 통합 테스트."""

    def test_ingest_to_render(
        self,
        pipeline_root,
        definition_path,
        docx_template,
        xlsx_template,
        sample_packet,
    ):
        """Ingest → Validate → SSOT → Render 전체 흐름."""
        jobs_dir = pipeline_root / "jobs"
        config = {
            "pipeline": {"hash_fields": ["wo_no", "line", "part_no", "lot", "result"]},
            "lock": {"timeout": 5},
        }

        # 1. Job ID 생성
        job_id = generate_job_id(sample_packet)
        job_dir = jobs_dir / job_id
        job_dir.mkdir(parents=True)

        # 2. Job 폴더 구조 생성 (SSOT)
        ensure_job_json(job_dir, sample_packet, config, generate_job_id)
        assert (job_dir / "job.json").exists()

        # 3. Run 로그 생성
        run_log = create_run_log(job_id)
        run_id = run_log.run_id
        assert run_id is not None

        # 4. 검증
        validation_service = ValidationService(definition_path)
        validation_result = validation_service.validate(sample_packet)
        assert validation_result.valid is True

        # 5. 해시 계산
        packet_hash = compute_packet_hash(sample_packet, config, definition_path)
        run_log.packet_hash = packet_hash

        # 6. DOCX 렌더링
        docx_output = job_dir / "deliverables" / "report.docx"
        docx_renderer = DocxRenderer(docx_template)
        docx_renderer.render(sample_packet, docx_output)
        assert docx_output.exists()

        # 7. XLSX 렌더링
        xlsx_output = job_dir / "deliverables" / "measurements.xlsx"
        xlsx_manifest = {
            "xlsx_mappings": {
                "named_ranges": {
                    "wo_no": "WO_NO",
                    "line": "LINE",
                    "result": "RESULT",
                },
                "cell_addresses": {},
            },
        }
        xlsx_renderer = ExcelRenderer(xlsx_template, xlsx_manifest)
        xlsx_renderer.render(sample_packet, xlsx_output)
        assert xlsx_output.exists()

        # 8. Run 완료 (complete_run_log는 in-place 수정, None 반환)
        complete_run_log(run_log, success=True, packet_hash=packet_hash)
        assert run_log.result == "success"
        assert run_log.finished_at is not None

        # 9. Run 로그 저장 (파일명: run_{run_id}.json)
        runs_dir = job_dir / "runs"
        save_run_log(run_log, runs_dir)
        assert (runs_dir / f"run_{run_id}.json").exists()

    def test_pipeline_creates_all_deliverables(
        self,
        pipeline_root,
        definition_path,
        docx_template,
        xlsx_template,
        sample_packet,
    ):
        """파이프라인이 모든 산출물 생성."""
        jobs_dir = pipeline_root / "jobs"
        job_id = generate_job_id(sample_packet)
        job_dir = jobs_dir / job_id
        job_dir.mkdir(parents=True)

        config = {"lock": {"timeout": 5}}

        # 전체 파이프라인 실행
        ensure_job_json(job_dir, sample_packet, config, generate_job_id)

        # 렌더링
        docx_output = job_dir / "deliverables" / "report.docx"
        xlsx_output = job_dir / "deliverables" / "measurements.xlsx"

        DocxRenderer(docx_template).render(sample_packet, docx_output)
        ExcelRenderer(
            xlsx_template,
            {
                "xlsx_mappings": {
                    "named_ranges": {"wo_no": "WO_NO"},
                    "cell_addresses": {},
                },
            },
        ).render(sample_packet, xlsx_output)

        # 검증
        deliverables = list((job_dir / "deliverables").iterdir())
        assert len(deliverables) == 2
        assert any(f.suffix == ".docx" for f in deliverables)
        assert any(f.suffix == ".xlsx" for f in deliverables)


# =============================================================================
# 7. 오류 상황 테스트
# =============================================================================


class TestPipelineErrorScenarios:
    """파이프라인 오류 상황 테스트."""

    def test_validation_failure_stops_pipeline(self, definition_path):
        """검증 실패 시 파이프라인 중단."""
        service = ValidationService(definition_path)

        # 불완전한 데이터
        incomplete = {"wo_no": "WO-001"}

        result = service.validate(incomplete)

        assert result.valid is False
        # 파이프라인은 여기서 중단되어야 함

    def test_missing_template_raises_error(self, tmp_path, sample_packet):
        """템플릿 없으면 에러."""
        from src.domain.errors import PolicyRejectError

        nonexistent = tmp_path / "nonexistent.docx"

        with pytest.raises(PolicyRejectError):
            DocxRenderer(nonexistent)
