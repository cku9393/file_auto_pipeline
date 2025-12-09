"""
test_word.py - Word (DOCX) 렌더러 테스트

테스트 대상:
- DocxRenderer: 템플릿 로드, 렌더링, placeholder 추출
- render_docx: 간편 함수
"""

from pathlib import Path

import pytest
from docx import Document

from src.domain.errors import ErrorCodes, PolicyRejectError
from src.render.word import DocxRenderer, render_docx

# =============================================================================
# Fixtures
# =============================================================================


@pytest.fixture
def simple_docx_template(tmp_path: Path) -> Path:
    """
    간단한 DOCX 템플릿 생성.

    placeholder: {{wo_no}}, {{line}}, {{result}}
    """
    template_path = tmp_path / "template.docx"

    doc = Document()
    doc.add_heading("검사 성적서", 0)
    doc.add_paragraph("작업번호: {{wo_no}}")
    doc.add_paragraph("라인: {{line}}")
    doc.add_paragraph("결과: {{result}}")
    doc.save(template_path)

    return template_path


@pytest.fixture
def template_with_measurements(tmp_path: Path) -> Path:
    """
    측정 데이터를 포함하는 DOCX 템플릿.

    placeholder: {{measurements}}
    """
    template_path = tmp_path / "template_measurements.docx"

    doc = Document()
    doc.add_heading("검사 성적서", 0)
    doc.add_paragraph("작업번호: {{wo_no}}")

    # 측정 테이블 (Jinja2 for loop)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "항목"
    hdr_cells[1].text = "규격"
    hdr_cells[2].text = "측정값"
    hdr_cells[3].text = "결과"

    # docxtpl에서는 테이블 내 for loop을 별도 처리해야 하므로
    # 여기서는 간단히 텍스트로 placeholder만 추가
    doc.add_paragraph("측정 데이터: {{measurements}}")

    doc.save(template_path)

    return template_path


@pytest.fixture
def sample_render_data() -> dict:
    """렌더링 테스트용 데이터."""
    return {
        "wo_no": "WO-001",
        "line": "L1",
        "part_no": "PART-A",
        "lot": "LOT-001",
        "result": "PASS",
        "inspector": "홍길동",
        "date": "2024-01-15",
        "remark": "테스트 비고",
        "measurements": [
            {"item": "길이", "spec": "10±0.1", "measured": "10.05", "result": "PASS"},
            {"item": "폭", "spec": "5±0.1", "measured": "5.02", "result": "PASS"},
        ],
    }


# =============================================================================
# DocxRenderer 초기화 테스트
# =============================================================================


class TestDocxRendererInit:
    """DocxRenderer 초기화 테스트."""

    def test_init_with_valid_template(self, simple_docx_template: Path):
        """유효한 템플릿으로 초기화."""
        renderer = DocxRenderer(simple_docx_template)

        assert renderer.template_path == simple_docx_template
        assert renderer._doc is None  # lazy loading

    def test_init_with_nonexistent_template(self, tmp_path: Path):
        """존재하지 않는 템플릿 → PolicyRejectError."""
        nonexistent = tmp_path / "nonexistent.docx"

        with pytest.raises(PolicyRejectError) as exc_info:
            DocxRenderer(nonexistent)

        assert exc_info.value.code == ErrorCodes.TEMPLATE_NOT_FOUND
        assert "path" in exc_info.value.context


# =============================================================================
# DocxRenderer.render 테스트
# =============================================================================


class TestDocxRendererRender:
    """DocxRenderer.render 테스트."""

    def test_basic_render(
        self,
        simple_docx_template: Path,
        sample_render_data: dict,
        tmp_path: Path,
    ):
        """기본 렌더링."""
        renderer = DocxRenderer(simple_docx_template)
        output_path = tmp_path / "output.docx"

        result = renderer.render(sample_render_data, output_path)

        assert result == output_path
        assert output_path.exists()

        # 렌더링 결과 확인
        doc = Document(output_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "WO-001" in full_text
        assert "L1" in full_text
        assert "PASS" in full_text

    def test_creates_output_directory(
        self,
        simple_docx_template: Path,
        sample_render_data: dict,
        tmp_path: Path,
    ):
        """출력 디렉터리 자동 생성."""
        renderer = DocxRenderer(simple_docx_template)
        output_path = tmp_path / "nested" / "dir" / "output.docx"

        renderer.render(sample_render_data, output_path)

        assert output_path.exists()
        assert output_path.parent.exists()

    def test_render_with_korean(
        self,
        simple_docx_template: Path,
        tmp_path: Path,
    ):
        """한글 데이터 렌더링."""
        renderer = DocxRenderer(simple_docx_template)
        output_path = tmp_path / "output.docx"

        data = {
            "wo_no": "작업-001",
            "line": "라인A",
            "result": "합격",
        }

        renderer.render(data, output_path)

        doc = Document(output_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "작업-001" in full_text
        assert "라인A" in full_text
        assert "합격" in full_text

    def test_render_with_measurements(
        self,
        template_with_measurements: Path,
        sample_render_data: dict,
        tmp_path: Path,
    ):
        """측정 데이터 포함 렌더링."""
        renderer = DocxRenderer(template_with_measurements)
        output_path = tmp_path / "output.docx"

        renderer.render(sample_render_data, output_path)

        assert output_path.exists()

    def test_render_with_missing_fields(
        self,
        simple_docx_template: Path,
        tmp_path: Path,
    ):
        """누락 필드는 빈 문자열로 처리."""
        renderer = DocxRenderer(simple_docx_template)
        output_path = tmp_path / "output.docx"

        # 일부 필드만 제공
        data = {"wo_no": "WO-001"}

        renderer.render(data, output_path)

        doc = Document(output_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "WO-001" in full_text
        # 누락 필드는 빈 문자열 (에러 아님)

    def test_render_includes_generated_at(
        self,
        simple_docx_template: Path,
        sample_render_data: dict,
        tmp_path: Path,
    ):
        """generated_at 메타데이터 추가 (컨텍스트에)."""
        renderer = DocxRenderer(simple_docx_template)
        output_path = tmp_path / "output.docx"

        # _build_context 내부에서 generated_at이 추가됨
        # 직접 확인하려면 템플릿에 {{generated_at}} 필요
        # 여기서는 에러 없이 렌더링되는지만 확인
        renderer.render(sample_render_data, output_path)
        assert output_path.exists()


# =============================================================================
# DocxRenderer.get_placeholders 테스트
# =============================================================================


class TestDocxRendererGetPlaceholders:
    """DocxRenderer.get_placeholders 테스트."""

    def test_extracts_placeholders(self, simple_docx_template: Path):
        """placeholder 목록 추출."""
        renderer = DocxRenderer(simple_docx_template)

        placeholders = renderer.get_placeholders()

        # {{wo_no}}, {{line}}, {{result}} 포함
        assert "wo_no" in placeholders
        assert "line" in placeholders
        assert "result" in placeholders

    def test_empty_template_returns_empty_list(self, tmp_path: Path):
        """placeholder 없는 템플릿."""
        template_path = tmp_path / "empty.docx"
        doc = Document()
        doc.add_paragraph("No placeholders here")
        doc.save(template_path)

        renderer = DocxRenderer(template_path)
        placeholders = renderer.get_placeholders()

        assert placeholders == []


# =============================================================================
# render_docx 간편 함수 테스트
# =============================================================================


class TestRenderDocx:
    """render_docx 간편 함수 테스트."""

    def test_basic_usage(
        self,
        simple_docx_template: Path,
        sample_render_data: dict,
        tmp_path: Path,
    ):
        """기본 사용."""
        output_path = tmp_path / "output.docx"

        result = render_docx(
            simple_docx_template,
            sample_render_data,
            output_path,
        )

        assert result == output_path
        assert output_path.exists()

    def test_with_photos(
        self,
        simple_docx_template: Path,
        sample_render_data: dict,
        tmp_path: Path,
    ):
        """사진 포함 렌더링 (사진 경로가 유효하면)."""
        output_path = tmp_path / "output.docx"

        # 가짜 사진 파일 생성
        photo_path = tmp_path / "photo.jpg"
        # 실제로는 유효한 이미지가 필요하지만,
        # 템플릿에 {{photo_overview}}가 없으면 무시됨
        photo_path.write_bytes(b"fake jpg content")

        result = render_docx(
            simple_docx_template,
            sample_render_data,
            output_path,
            photos={"overview": photo_path},
        )

        assert result == output_path
